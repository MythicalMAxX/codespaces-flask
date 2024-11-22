import os
from flask import Flask, request, jsonify, send_from_directory, render_template
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from langchain.chains import ConversationalRetrievalChain
from langchain_community.chat_models import ChatOpenAI
from langchain_community.document_loaders import DirectoryLoader
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.indexes import VectorstoreIndexCreator
from langchain.schema import Document as LangchainDocument
from dotenv import load_dotenv

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///chat.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db = SQLAlchemy(app)
migrate = Migrate(app, db)

# Load environment variables from .env file
load_dotenv()

os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

# Enable to save to disk & reuse the model (for repeated queries on the same data)
PERSIST = False

embedding_function = OpenAIEmbeddings()

if PERSIST and os.path.exists("persist"):
    print("Reusing index...\n")
    vectorstore = Chroma(persist_directory="persist", embedding_function=embedding_function)
    index = VectorStoreIndexWrapper(vectorstore=vectorstore)
else:
    loader = DirectoryLoader("Data/")
    if PERSIST:
        index = VectorstoreIndexCreator(vectorstore_kwargs={"persist_directory":"persist"}, embedding=embedding_function).from_loaders([loader])
    else:
        index = VectorstoreIndexCreator(embedding=embedding_function).from_loaders([loader])

chain = ConversationalRetrievalChain.from_llm(
    llm=ChatOpenAI(model="gpt-3.5-turbo"),
    retriever=index.vectorstore.as_retriever(search_kwargs={"k": 1}),
)

chat_history = []

class ChatHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_message = db.Column(db.String, nullable=False)
    bot_response = db.Column(db.String, nullable=False)

class FileData(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    file_name = db.Column(db.String, nullable=False)
    file_content = db.Column(db.LargeBinary, nullable=False)

def save_chat_to_db(user_message, bot_response):
    chat = ChatHistory(user_message=user_message, bot_response=bot_response)
    db.session.add(chat)
    db.session.commit()

def load_chat_history():
    return ChatHistory.query.all()

def save_file_data_to_db(file_name, file_content):
    file_data = FileData(file_name=file_name, file_content=file_content)
    db.session.add(file_data)
    db.session.commit()

def load_file_data():
    return FileData.query.all()

def parse_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def parse_docx(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text

def parse_txt(file_path):
    with open(file_path, 'r') as file:
        text = file.read()
    return text

@app.route('/chat', methods=['POST'])
def chat():
    global chat_history
    user_message = request.json.get('message')
    if not user_message:
        return jsonify({"error": "No message provided"}), 400

    file_data = load_file_data()
    combined_text = " ".join([data.file_content.decode('utf-8') if isinstance(data.file_content, bytes) else data.file_content for data in file_data])
    documents = [LangchainDocument(page_content=combined_text)]

    # Create a new index from the documents
    index = VectorstoreIndexCreator(embedding=embedding_function).from_documents(documents)

    chain = ConversationalRetrievalChain.from_llm(
        llm=ChatOpenAI(model="gpt-3.5-turbo"),
        retriever=index.vectorstore.as_retriever(search_kwargs={"k": 1}),
    )

    result = chain({"question": user_message, "chat_history": chat_history})
    response_text = result['answer']
    chat_history.append((user_message, response_text))
    save_chat_to_db(user_message, response_text)

    return jsonify({"response": response_text})

@app.route('/upload', methods=['POST'])
def upload_files():
    try:
        if 'examSchedule' not in request.files or 'syllabus' not in request.files:
            return jsonify({"error": "No file part"}), 400

        exam_schedule = request.files['examSchedule']
        syllabus = request.files['syllabus']

        if exam_schedule.filename == '' or syllabus.filename == '':
            return jsonify({"error": "No selected file"}), 400

        exam_schedule_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(exam_schedule.filename))
        syllabus_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(syllabus.filename))

        exam_schedule.save(exam_schedule_path)
        syllabus.save(syllabus_path)

        # Parse files and save content to database
        for file_path in [exam_schedule_path, syllabus_path]:
            if file_path.endswith('.pdf'):
                file_content = parse_pdf(file_path).encode('utf-8')
            elif file_path.endswith('.docx'):
                file_content = parse_docx(file_path).encode('utf-8')
            elif file_path.endswith('.txt'):
                file_content = parse_txt(file_path).encode('utf-8')
            else:
                continue
            save_file_data_to_db(os.path.basename(file_path), file_content)

        return jsonify({"success": "Files uploaded and parsed successfully"}), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/<path:path>')
def serve_static(path):
    return send_from_directory('static', path)

if __name__ == "__main__":
    app.run(debug=True)